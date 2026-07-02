"""
文档生成服务：Jinja2 渲染 Markdown 模板 + python-docx 导出 Word。
"""
import os
import uuid
from jinja2 import Template
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from app.config import settings
from app.database import get_connection


def get_template_by_id(template_id: int) -> dict:
    conn = get_connection()
    row = conn.execute("SELECT * FROM templates WHERE id = ?", (template_id,)).fetchone()
    conn.close()
    return dict(row) if row else {}


def get_all_templates(template_type: str = None) -> list[dict]:
    conn = get_connection()
    if template_type:
        rows = conn.execute("SELECT * FROM templates WHERE template_type = ? ORDER BY id", (template_type,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM templates ORDER BY id").fetchall()
    conn.close()
    return [dict(r) for r in rows]


def create_template(name: str, template_type: str, content: str, variables: str = "") -> int:
    conn = get_connection()
    cursor = conn.execute(
        "INSERT INTO templates (name, template_type, content, variables) VALUES (?, ?, ?, ?)",
        (name, template_type, content, variables)
    )
    conn.commit()
    tid = cursor.lastrowid
    conn.close()
    return tid


def delete_template(template_id: int):
    conn = get_connection()
    conn.execute("DELETE FROM templates WHERE id = ?", (template_id,))
    conn.commit()
    conn.close()


def render_document(parsed_spec: dict, selected_clauses: list[dict],
                    risk_report: dict, template_id: int = 1) -> str:
    """使用 Jinja2 渲染模板生成文档文本。"""
    template_data = get_template_by_id(template_id)
    if not template_data:
        # 默认模板
        template_content = _default_template()
    else:
        template_content = template_data["content"]

    tpl = Template(template_content)
    context = {
        **parsed_spec,
        "selected_clauses": selected_clauses,
        "risk_report": risk_report,
        "risks": risk_report.get("risks", []),
        "overall_level": risk_report.get("overall_level", "low"),
    }
    # 确保列表字段有默认值
    for key in ["qualification", "technical_specs", "evaluation_factors"]:
        if key not in context or context[key] is None:
            context[key] = []

    return tpl.render(**context)


def export_to_docx(rendered_text: str, project_name: str) -> str:
    """将渲染文本导出为专业格式的 .docx 文件。"""
    doc = Document()

    # ── 全局默认样式 ──
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    rpr = style.element.rPr
    if rpr is None:
        from docx.oxml import OxmlElement
        rpr = OxmlElement('w:rPr')
        style.element.append(rpr)
    from docx.oxml.ns import qn
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 解析内容 ──
    lines = rendered_text.strip().split('\n')
    in_table = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 表格行（包含 | 的行）
        if '|' in stripped and stripped.startswith('|'):
            # 跳过表头分隔线
            if all(c in '|-: ' for c in stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                table = doc.add_table(rows=1, cols=len(cells), style='Light Grid Accent 1')
                table.autofit = True
                in_table = True
                # 表头加粗
                for i, cell_text in enumerate(cells):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)
            else:
                row = table.add_row()
                for i, cell_text in enumerate(cells):
                    row.cells[i].text = cell_text
                    for p in row.cells[i].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
            continue
        else:
            in_table = False

        # 一级标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            h = doc.add_heading(stripped[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.size = Pt(22)
        # 二级标题
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        # 三级标题
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        # 引用
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
        # 分隔线
        elif stripped.startswith('---'):
            doc.add_paragraph('─' * 50)
        else:
            doc.add_paragraph(stripped)

    # ── 页脚 ──
    footer = section.footer
    footer.paragraphs[0].text = f"ProcureGen AI 自动生成 | {project_name or '采购文件'}"
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.paragraphs[0].runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 保存 ──
    filename = f"{project_name or '采购文件'}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(settings.EXPORT_DIR, filename)
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    doc.save(filepath)
    return filepath, filename


def _default_template() -> str:
    """默认招标书模板。"""
    return """# {{project_name or '采购项目'}} 招标文件

---

## 第一章 采购公告

### 1.1 项目概况

| 项目 | 内容 |
|------|------|
| 项目名称 | {{project_name or '（待填写）'}} |
| 采购类型 | {{purchase_type or '（待填写）'}} |
| 采购方式 | {{procurement_method or '公开招标'}} |
| 预算金额 | {{budget or '（待填写）'}} 元 |
| 交付周期 | {{delivery_period or '（待填写）'}} |

---

## 第二章 供应商须知

### 2.1 供应商资格要求

{% if qualification %}
{% for item in qualification %}
- {{item}}
{% endfor %}
{% else %}
- （暂无资格要求，请补充）
{% endif %}

---

## 第三章 采购需求

### 3.1 技术参数

{% if technical_specs %}
{% for item in technical_specs %}
- {{item}}
{% endfor %}
{% else %}
- （暂无技术参数，请补充）
{% endif %}

### 3.2 交付与验收

- 交付周期：{{delivery_period or '（待填写）'}}
- 验收标准：{{acceptance_criteria or '（待填写）'}}
- 质保期限：{{warranty_period or '（待填写）'}}

---

## 第四章 评标办法

### 4.1 评分因素

{% if evaluation_factors %}
{% for item in evaluation_factors %}
- {{item}}
{% endfor %}
{% else %}
- 价格
- 技术
- 商务
- 售后服务
{% endif %}

---

## 第五章 合同主要条款

{% if selected_clauses %}
{% for clause in selected_clauses %}
### {{clause.title}}

{{clause.content}}

> 推荐理由：{{clause.reason}}

{% endfor %}
{% else %}
（暂无推荐条款）
{% endif %}

### 付款方式

{{payment_terms or '（待填写）'}}

---

## 第六章 风险提示附录

> 风险等级：{{overall_level or 'low'}}

{% if risks %}
{% for risk in risks %}
### {{risk.level | upper}} 风险：{{risk.type}}

- **说明**：{{risk.message}}
- **建议**：{{risk.suggestion}}

{% endfor %}
{% else %}
暂无风险提示。
{% endif %}

---

*本文档由 ProcureGen AI 自动生成，仅供采购参考，最终版本请人工复核确认。*
"""
